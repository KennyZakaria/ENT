import os
import json
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
import aiofiles
from PyPDF2 import PdfReader
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
from utils.config import config
from models import Document

class DocumentService:
    def __init__(self):
        self.documents_file = config.DOCUMENTS_JSON
        
    async def load_documents(self) -> List[Document]:
        """Charge la liste des documents depuis le fichier JSON"""
        try:
            if os.path.exists(self.documents_file):
                async with aiofiles.open(self.documents_file, 'r', encoding='utf-8') as f:
                    content = await f.read()
                    data = json.loads(content)
                    return [Document(**doc) for doc in data]
            return []
        except Exception:
            return []
    
    async def save_documents(self, documents: List[Document]):
        """Sauvegarde la liste des documents dans le fichier JSON"""
        data = [doc.dict() for doc in documents]
        async with aiofiles.open(self.documents_file, 'w', encoding='utf-8') as f:
            await f.write(json.dumps(data, default=str, indent=2))
    
    async def add_document(self, filename: str, file_content: bytes) -> Document:
        """Ajoute un nouveau document"""
        doc_id = str(uuid.uuid4())
        file_path = os.path.join(config.UPLOAD_DIR, f"{doc_id}_{filename}")
        
        # Sauvegarder le fichier
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Extraire le texte
        text_content = await self._extract_text(file_path, filename)
        
        # Sauvegarder le contenu texte
        text_path = os.path.join(config.UPLOAD_DIR, f"{doc_id}.txt")
        async with aiofiles.open(text_path, 'w', encoding='utf-8') as f:
            await f.write(text_content)
        
        # Compter les chunks (approximatif)
        chunk_count = max(1, len(text_content) // config.CHUNK_SIZE)
        
        # Créer l'objet Document
        document = Document(
            id=doc_id,
            filename=filename,
            file_size=len(file_content),
            ingested_at=datetime.now(),
            chunk_count=chunk_count
        )
        
        # Ajouter à la liste
        documents = await self.load_documents()
        documents.append(document)
        await self.save_documents(documents)
        
        return document
    
    async def delete_document(self, doc_id: str) -> bool:
        """Supprime un document"""
        documents = await self.load_documents()
        documents = [doc for doc in documents if doc.id != doc_id]
        await self.save_documents(documents)
        
        # Supprimer les fichiers physiques
        try:
            for file in os.listdir(config.UPLOAD_DIR):
                if file.startswith(doc_id):
                    os.remove(os.path.join(config.UPLOAD_DIR, file))
        except:
            pass
        
        return True
    
    async def get_document_content(self, doc_id: str) -> Optional[str]:
        """Récupère le contenu d'un document"""
        text_path = os.path.join(config.UPLOAD_DIR, f"{doc_id}.txt")
        if os.path.exists(text_path):
            async with aiofiles.open(text_path, 'r', encoding='utf-8') as f:
                return await f.read()
        return None
    
    async def _extract_text(self, file_path: str, filename: str) -> str:
        """Extrait le texte d'un fichier"""
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'txt':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        
        elif file_ext == 'pdf':
            return await self._extract_pdf_text(file_path)
        
        elif file_ext == 'docx':
            if DocxDocument is None:
                raise ValueError("python-docx n'est pas installé. Installez-le avec: pip install python-docx")
            return await self._extract_docx_text(file_path)
        
        else:
            raise ValueError(f"Format de fichier non supporté: {file_ext}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extrait le texte d'un PDF"""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Erreur lors de l'extraction PDF: {str(e)}")
        return text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extrait le texte d'un document Word"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Erreur lors de l'extraction DOCX: {str(e)}")

# Instance globale
document_service = DocumentService()